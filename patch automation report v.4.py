# -*- coding: utf-8 -*-
"""
Created on Thu Jun 30 11:36:29 2022

@author: ctagle
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Jun 30 09:01:59 2022

@author: ctagle
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Jun  2 12:05:09 2022

@author: ctagle
"""

import os
print(os.getcwd())
import pandas as pd
from openpyxl import load_workbook
import docx
from docx import Document
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.filedialog as fd
from datetime import date
import os

tdd= str(date.today())
CurrentMonth = date.today().month
def RangeOFDates ():
    if CurrentMonth==1:
        DR=str('Jan 1 - Jan 31')
    elif CurrentMonth==2:
        DR=str('Feb 1 - Feb 28')
    elif CurrentMonth==3:
        DR=str('Mar 1 - Mar 30')
    elif CurrentMonth==4:
        DR=str('Apr 1 - Apr 31')
    elif CurrentMonth==5:
        DR=str('May 1 - May 30')
    elif CurrentMonth==6:
        DR=str('June 1 - June 31')
    elif CurrentMonth==7:
        DR=str('July 1 - July 30')
    elif CurrentMonth==8:
        DR=str('Aug 1 - Aug 31')
    elif CurrentMonth==9:
        DR=str('Sept 1 - Sept 30')
    elif CurrentMonth==10:
        DR=str('Oct 1 - Oct 31')
    elif CurrentMonth==11:
        DR=str('Nov 1 - Nov 30')
    elif CurrentMonth==12:
        DR=str('Dec 1 - Dec 31')
    return DR
    
def generatereport ():
    read_file = pd.read_csv ( filenames[0])
    read_file.to_excel (r'C:\Temp\test1 patch automation.xlsx', index = None, header=True)
    
    wb = load_workbook(r'C:\Temp\test1 patch automation.xlsx')
    ws = wb.active
    
    patchlist=[]
    devicename=[]
    device_name_col = 16
    patch_name_col = 13
    fclient = ws['A8'].value
    client = str(fclient)
    
    for row in ws.iter_rows(min_row=18, values_only=True):
        if (row[patch_name_col] != None):
            patchlist.append(row[patch_name_col])
            devicename.append(row[device_name_col])
            
    
            
    document = Document()
    header = document.sections[0].header
    hdrr = header.paragraphs[0]
    logo = hdrr.add_run()
    logo.text = ('\t' + '\t')
    logo.add_picture('Deltra Logo.png')
    
    foot = document.sections[0].footer
    footr = foot.paragraphs[0]
    footer = footr.add_run()
    footer.add_picture('Deltra Footer.png', width=docx.shared.Inches(5.80), height=docx.shared.Inches(.6))
    
    document.add_heading('Patch Remediation Report', 0)
    p = document.add_paragraph('This document serves to inform and provide recommendations by Deltra Systems LLC for the following list of patches that are not supported under our current patch management program, and it is advised that each of the below items be remediated by the client as soon as possible in order to ensure that the environment continues to operate within industry') 
    p.add_run('best practices.')
    p1 = document.add_paragraph('Failure to apply the following recommendations may result in security vulnerabilities and system instabilities on each of the following devices.', style = 'Heading 1')
    
    p2 = document.add_heading('', 0)
    
    CN = document.add_paragraph().add_run('Client Name: ' + client, style = 'Strong')
    RD = document.add_paragraph().add_run('Report Date: ' + tdd, style = 'Strong')
    RN = document.add_paragraph().add_run("Reviewer's Name: " + Reviewer.get(), style = 'Strong')
    PD = document.add_paragraph().add_run('Patch Date Range: ' + RangeOFDates(), style = 'Strong')
    
    table = document.add_table(rows=len(patchlist)+1, cols=3, style='Table Grid')
    tb_cells = table.rows[0].cells
    tb_cells[0].text = 'Device Name'
    tb_cells[1].text = 'Patch Name'
    tb_cells[2].text = 'Remediation Action'
    
    for i in range(len(patchlist)):
        table.rows[i+1].cells[1].text = patchlist[i]
        table.rows[i+1].cells[0].text = devicename[i]
        
    document.save(client+' PatchReport '+tdd+'.docx')
    wb.remove(r'C:\Temp\test1 patch automation.xlsx')
    

root = tk.Tk()

root.title('Patch Remediation Report Generator (PRRG)')

root_width = 600
root_height = 400

screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

center_x = int(screen_width/2 - root_width/2)
center_y = int(screen_height/2 - root_height/2)

root.geometry(f'{root_width}x{root_height}+{center_x}+{center_y}')

root.columnconfigure(0, weight=1)
root.columnconfigure(1, weight=3)

root.attributes('-topmost')


filenames = str()

def select_files():
    global filenames
    filetypes = (('all files', '*.*'),
                 ('CSV', '*.csv'))
    filenames = fd.askopenfilenames(
        title='Select Files', 
        initialdir='/', 
        filetypes = filetypes
        )
    return filenames
    
open_button = ttk.Button(
        root,
        text='Select File',
        command = select_files
        )
open_button.grid(
    column=1,
    row=1,
    padx=5, 
    pady=5
    )
open_label = ttk.Label(root, text='Select .CSV to Generate:')
open_label.grid(
    column=0,
    row=1,
    padx=0,
    pady=5
    )
reviewer = str()
Reviewer = ttk.Entry(root , textvariable=reviewer, )
Reviewer.grid(
    column=1,
    row=2,
    padx=0,
    pady=5
    )
Reviewer_label=ttk.Label(root, text='Reviewers Name:')
Reviewer_label.grid(
    column=0,
    row=2,
    padx=0,
    pady=5
    )
generate_btn = ttk.Button(root, text='Generate Report', command=generatereport )
generate_btn.grid(
    columnspan=2,
    row=6,
    padx=0,
    pady=160
    )



    



root.mainloop()
